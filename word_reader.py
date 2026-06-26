from io import BytesIO
from docx import Document
import pandas as pd

from config import (
    MAIN_TABLE_INDEX,
    WORD_FIELDS,
    DUPLICATE_COLUMNS,
    MILESTONE_HEADER_ROW,
    MILESTONE_PARENT_ROW,
    MILESTONE_COLUMNS,
    MILESTONE_PARENT_COLUMN,
    MILESTONE_TABLE_INDEX
)

class WordReader:

    def __init__(self, file):
        # Accept either a bytes object or a file-like object (UploadedFile / BytesIO)
        file_bytes = None
        try:
            file_bytes = file.read() if hasattr(file, "read") else file
            if isinstance(file_bytes, memoryview):
                file_bytes = file_bytes.tobytes()
            # Wrap into BytesIO for python-docx
            bio = BytesIO(file_bytes)
            try:
                self.doc = Document(bio)
            except Exception as e:
                # Give a clear message so caller can see whether it's a format/problem
                raise ValueError(f"python-docx failed to open the file as .docx: {e}") from e

            # Validate expected table exists
            if not getattr(self.doc, "tables", None) or len(self.doc.tables) <= MAIN_TABLE_INDEX:
                raise ValueError(f"Expected table index {MAIN_TABLE_INDEX} not found in the uploaded document. Document contains {len(self.doc.tables) if getattr(self.doc, 'tables', None) is not None else 0} tables.")
            self.table = self.doc.tables[MAIN_TABLE_INDEX]
        except Exception:
            # Re-raise so calling code can display traceback
            raise

    def read_fields(self):
        data = {}
        for field, (row, col) in WORD_FIELDS.items():
            value = self.table.rows[row].cells[col].text.strip()
            data[field] = value
        return data

    def read_milestones(self):
        parent_cell = self.table.rows[MILESTONE_PARENT_ROW].cells[MILESTONE_PARENT_COLUMN]
        nested_table = parent_cell.tables[MILESTONE_TABLE_INDEX]
        milestones = []

        for row in nested_table.rows[MILESTONE_HEADER_ROW+1:]:
            values = [cell.text.strip() for cell in row.cells]
            if not any(values):
                continue

            milestone = {
                "name": values[MILESTONE_COLUMNS["name"]],
                "date": values[MILESTONE_COLUMNS["date"]],
                "monthly": values[MILESTONE_COLUMNS["monthly"]],
                "quality": values[MILESTONE_COLUMNS["quality"]],
                "invoice": values[MILESTONE_COLUMNS["invoice"]]
            }
            milestones.append(milestone)

        return pd.DataFrame(milestones)


def read_word(file):
    # Let WordReader raise meaningful errors which app.py can display
    reader = WordReader(file)
    return {
        "fields": reader.read_fields(),
        "milestones": reader.read_milestones()
    }
