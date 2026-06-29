from io import BytesIO
from copy import deepcopy
from docx import Document
import pandas as pd

from config import (
    MAIN_TABLE_INDEX,
    WORD_FIELDS,
    DUPLICATE_COLUMNS,
    MILESTONE_COLUMNS,
    MILESTONE_HEADER_ROW,
    MILESTONE_PARENT_COLUMN,
    MILESTONE_PARENT_ROW,
    MILESTONE_TABLE_INDEX
)

class WordUpdater:

    def __init__(self, file):
        file_bytes = file.read() if hasattr(file, "read") else file
        if isinstance(file_bytes, memoryview):
            file_bytes = file_bytes.tobytes()
        self.doc = Document(BytesIO(file_bytes))
        self.table = self.doc.tables[MAIN_TABLE_INDEX]

    def update_fields(self, field_values):
        for field_name, value in field_values.items():
            if field_name not in WORD_FIELDS:
                continue
            
            row, _ = WORD_FIELDS[field_name]
            if field_name == "Total Cost":
                clean_value = str(value).replace("€", "").strip()
                display_value = f"€ {clean_value}"
            else:
                display_value = str(value)
                
            for col in DUPLICATE_COLUMNS:
                self.table.rows[row].cells[col].text = display_value

    def get_milestone_table(self):
        parent_cell = self.table.rows[MILESTONE_PARENT_ROW].cells[MILESTONE_PARENT_COLUMN]
        return parent_cell.tables[MILESTONE_TABLE_INDEX]

    def clear_milestone(self):
        table = self.get_milestone_table()
        while len(table.rows) > MILESTONE_HEADER_ROW + 1:
            tr = table.rows[-1]._tr
            tr.getparent().remove(tr)

    def update_milestone(self, milestone_df):
        table = self.get_milestone_table()
        if len(table.rows) > MILESTONE_HEADER_ROW + 1:
            template_row = deepcopy(table.rows[MILESTONE_HEADER_ROW + 1]._tr)
        else:
            raise ValueError("Milestone table has no template data row")

        while len(table.rows) > MILESTONE_HEADER_ROW + 1:
            tr = table.rows[-1]._tr 
            tr.getparent().remove(tr)

        # Standard value cleaning
        def clean_val(val):
            if pd.isna(val) or val is None:
                return ""
            return str(val).strip()

        # Currency formatting utility for floats
        def clean_currency_val(val):
            if pd.isna(val) or val is None or str(val).strip() == "":
                return ""
            # Strip out any existing symbols or extra spaces
            raw_num = str(val).replace("€", "").replace(",", "").strip()
            try:
                # Convert to float and format to exactly 2 decimal places
                formatted_num = f"{float(raw_num):,.2f}"
                return f"€ {formatted_num}"
            except ValueError:
                return f"€ {raw_num}"

        for _, milestone in milestone_df.iterrows():
            new_row = deepcopy(template_row)
            table._tbl.append(new_row)
            row = table.rows[-1]
            
            row.cells[MILESTONE_COLUMNS["name"]].text = clean_val(milestone["name"])
            row.cells[MILESTONE_COLUMNS["date"]].text = clean_val(milestone["date"])
            
            # Explicit Financial Currency fields
            row.cells[MILESTONE_COLUMNS["monthly"]].text = clean_currency_val(milestone["monthly"])
            row.cells[MILESTONE_COLUMNS["quality"]].text = clean_currency_val(milestone["quality"])
            row.cells[MILESTONE_COLUMNS["invoice"]].text = clean_currency_val(milestone["invoice"])

    def update(self, field_values, milestone_df):
        self.update_fields(field_values)
        self.update_milestone(milestone_df)

    def save(self):
        output = BytesIO()
        self.doc.save(output)
        output.seek(0)
        return output

def update_word(uploaded_file, field_values, milestone_df):
    updater = WordUpdater(uploaded_file)
    updater.update(field_values, milestone_df)
    return updater.save()
