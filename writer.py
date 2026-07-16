import docx
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import re
from typing import Optional, Dict, List
from copy import deepcopy

def remove_row(table, row):
    """Removes a row from a table."""
    tbl = table._tbl
    tr = row._tr
    tbl.remove(tr)

def set_cell_font(cell, font_name="Arial", size_pt=10):
    """Formats all paragraphs and runs inside a cell to use a specific font."""
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.name = font_name
            run.font.size = Pt(size_pt)

def add_paragraph_with_run(cell, text, font_name="Arial", size_pt=10):
    """Helper to write text to a cell with correct font properties."""
    clear_cell(cell)
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    return p


def clear_cell(cell):
    """Remove all content from a table cell while preserving its properties."""
    tc = cell._tc
    for child in list(tc):
        if child.tag != qn("w:tcPr"):
            tc.remove(child)
    cell.add_paragraph("")

def ensure_table_gridspan_sync(table):
    """
    If a table header has an empty spacer cell at tc[1], merges it into tc[0] by:
    1. Setting gridSpan=2 on tc[0] so it spans 2 grid columns
    2. Merging tc[1]'s column width into tc[0]
    3. Removing tc[1]'s w:tc element from the w:tr
    This eliminates the blank intermediate column between ID and Description.
    """
    if not hasattr(table, 'rows') or not table.rows or len(table.rows) == 0:
        return
    # Only act on tables whose header tc[1] is an empty spacer
    hdr_tcs = list(table.rows[0]._tr.iterchildren(qn('w:tc')))
    if len(hdr_tcs) < 3:
        return
    hdr_tc1_text = "".join((p.text or "") for p in hdr_tcs[1].iterchildren(qn('w:p'))).strip()
    if hdr_tc1_text != '':
        return  # tc[1] has real content, not a spacer

    for row in table.rows:
        tcs = list(row._tr.iterchildren(qn('w:tc')))
        if len(tcs) < 3:
            continue
        tc0 = tcs[0]
        tc1 = tcs[1]

        # Skip if tc0 already has gridSpan >= 2
        tc0Pr = tc0.find(qn('w:tcPr'))
        gs_el = tc0Pr.find(qn('w:gridSpan')) if tc0Pr is not None else None
        if gs_el is not None and int(gs_el.get(qn('w:val'), '1')) >= 2:
            continue

        # Set gridSpan=2 on tc0
        if tc0Pr is None:
            tc0Pr = OxmlElement('w:tcPr')
            tc0.insert(0, tc0Pr)
        if gs_el is None:
            gs_el = OxmlElement('w:gridSpan')
            tc0Pr.append(gs_el)
        gs_el.set(qn('w:val'), '2')

        # Merge tc1's width into tc0
        tc0_w = tc0Pr.find(qn('w:tcW'))
        tc1Pr = tc1.find(qn('w:tcPr'))
        tc1_w = tc1Pr.find(qn('w:tcW')) if tc1Pr is not None else None
        if tc0_w is not None and tc1_w is not None:
            w0 = int(tc0_w.get(qn('w:w'), '0'))
            w1 = int(tc1_w.get(qn('w:w'), '0'))
            tc0_w.set(qn('w:w'), str(w0 + w1))

        # Remove tc1 from the row
        row._tr.remove(tc1)

def sync_and_clear_row_cells(table, row):
    """
    Clears old cell contents in a data row while preserving original gridSpan merged column layouts.
    """
    for cell in row.cells:
        clear_cell(cell)
    return row

def insert_row_after(table, row_idx, template_row_idx=None):
    """
    Surgically inserts a new row in a table after row_idx.
    Optionally specifies template_row_idx to copy data row formatting (borders, shading)
    instead of header row formatting.
    """
    if template_row_idx is None:
        template_row_idx = row_idx
    if template_row_idx >= len(table.rows):
        template_row_idx = len(table.rows) - 1

    template_row = table.rows[template_row_idx]
    template_tr = template_row._tr
    tr = deepcopy(template_tr)
    
    target_tr = table.rows[row_idx]._tr
    target_tr.addnext(tr)
    
    new_row = table.rows[row_idx + 1]
    sync_and_clear_row_cells(table, new_row)
    return new_row


def parse_requirement_order(req_id: str):
    """Return a sortable prefix/number pair for requirement IDs like URS-SF-100."""
    if not req_id:
        return "", None, 0

    normalized = req_id.strip().upper()
    match = re.match(r"^(.*?)(\d+)$", normalized)
    if not match:
        return normalized, None, 0

    prefix, number_text = match.groups()
    return prefix, int(number_text), len(number_text)


def find_insert_position(table, req_id: str):
    """Find the row index after which a new requirement should be inserted."""
    new_prefix, new_number, _ = parse_requirement_order(req_id)
    if new_number is None:
        return -1

    best_row_idx = -1
    best_number = -1

    for row_idx, row in enumerate(table.rows[1:], start=1):
        if len(row.cells) == 0:
            continue

        existing_id = row.cells[0].text.strip().upper()
        existing_prefix, existing_number, _ = parse_requirement_order(existing_id)

        if existing_number is None:
            continue

        if existing_prefix != new_prefix:
            continue

        if existing_number <= new_number and existing_number > best_number:
            best_number = existing_number
            best_row_idx = row_idx

    return best_row_idx


def get_requirement_row_number(row):
    """Return the numeric suffix for a requirement row, if present."""
    if len(row.cells) == 0:
        return None
    _, number, _ = parse_requirement_order(row.cells[0].text.strip().upper())
    return number


# F19: Extracted shared cell-population helper (was duplicated in MODIFY and INSERT blocks)
def get_table_column_indices(table):
    """
    Dynamically resolves column role indices based on table header label strings.
    Zero static column assumptions!
    """
    col_map = {
        "id": -1,
        "desc": -1,
        "impact": -1,
        "method": -1,
        "mappings": -1
    }
    if not hasattr(table, 'rows') or not table.rows or len(table.rows) == 0:
        return {
            "id": 0,
            "desc": 1,
            "impact": -1,
            "method": -1,
            "mappings": -1
        }

    ensure_table_gridspan_sync(table)

    header_cells = table.rows[0].cells
    mapping_candidates = []
    for idx, cell in enumerate(header_cells):
        txt = cell.text.lower().strip()
        if 'id' in txt or 'num' in txt or 'identifier' in txt or 'spec' in txt or txt == '#':
            if col_map["id"] == -1 and idx < 2:
                col_map["id"] = idx
        if 'description' in txt or 'specification' in txt:
            if 'id' not in txt and 'num' not in txt and 'identifier' not in txt:
                if col_map["desc"] == -1:
                    col_map["desc"] = idx
        if 'urs identifier' in txt or 'frs identifier' in txt or 'urs/frs' in txt or 'traceability' in txt:
            mapping_candidates.append(idx)
        elif 'mapping' in txt and 'id' not in txt and 'num' not in txt and 'identifier' not in txt:
            mapping_candidates.append(idx)
        if 'impact' in txt or 'patient safety' in txt or 'risk determination' in txt:
            if col_map["impact"] == -1:
                col_map["impact"] = idx
        if 'method' in txt or 'implementation' in txt or 'risk level' in txt:
            if col_map["method"] == -1:
                col_map["method"] = idx

    if col_map["mappings"] == -1 and mapping_candidates:
        col_map["mappings"] = mapping_candidates[-1]

    # Apply defaults if not resolved
    if col_map["id"] == -1:
        col_map["id"] = 0
    if col_map["desc"] == -1:
        col_map["desc"] = 1

    # Fallbacks based on cell count if mapping column wasn't explicitly named
    num_cols = len(header_cells)
    if col_map["mappings"] == -1:
        if num_cols == 3 and col_map["method"] != 2:
            col_map["mappings"] = 2
        elif num_cols >= 4 and col_map["method"] == 2:
            col_map["mappings"] = 3

    return col_map

def populate_row_cells(row, doc_type, data, table=None):
    """Populates row cells based on dynamically resolved column indices."""
    if table is None and hasattr(row, '_tr'):
        try:
            table = docx.table.Table(row._tr.getparent(), row._parent)
        except Exception:
            table = None

    col_map = get_table_column_indices(table) if table else {"id": 0, "desc": 1, "impact": 2 if doc_type == 'URS' else -1, "method": 3 if doc_type == 'URS' else (-1 if doc_type == 'FRS' else 2), "mappings": 2 if doc_type == 'FRS' else (3 if doc_type == 'DS' else -1)}
    
    desc_idx = col_map.get("desc", 1)
    if desc_idx < len(row.cells) and data.get("description") is not None:
        add_paragraph_with_run(row.cells[desc_idx], data.get("description", ""))

    impact_idx = col_map.get("impact", -1)
    if impact_idx != -1 and impact_idx < len(row.cells) and data.get("impact") is not None:
        add_paragraph_with_run(row.cells[impact_idx], data.get("impact"))

    method_idx = col_map.get("method", -1)
    if method_idx != -1 and method_idx < len(row.cells) and data.get("method") is not None:
        add_paragraph_with_run(row.cells[method_idx], data.get("method"))

    mappings_idx = col_map.get("mappings", -1)
    if mappings_idx != -1 and mappings_idx < len(row.cells) and data.get("mappings") is not None:
        add_paragraph_with_run(row.cells[mappings_idx], data.get("mappings"))


# F15: Find the correct requirement table by header content, not magic index
def find_requirement_table(doc, prefix):
    """Finds the table containing requirements matching the given prefix by scanning headers."""
    for table_idx, table in enumerate(doc.tables):
        if len(table.rows) == 0:
            continue
        # Check first column of data rows for prefix match
        for row in table.rows[1:]:  # Skip header row
            if len(row.cells) > 0:
                cell_text = row.cells[0].text.strip().upper()
                if cell_text.startswith(prefix.upper()):
                    return table_idx, table
    return None, None


# F16: Find the FIRST revision history table in the document
def find_revision_table(doc):
    """Finds the FIRST revision history table in the document by scanning headers forward."""
    revision_keywords = ['revision', 'rev', 'change history', 'document history']
    for table in doc.tables:
        if len(table.rows) == 0:
            continue
        header_text = " ".join(cell.text.strip().lower() for cell in table.rows[0].cells)
        if any(kw in header_text for kw in revision_keywords):
            return table
    # Fallback: first table if available
    return doc.tables[0] if len(doc.tables) > 0 else None


def rebuild_document_matrix(source_path: str, output_path: str, doc_type: str, changes: list, revision_log: str, revision_reason: str = "Updated as per user story modifications.", metadata_out: Optional[dict] = None) -> bool:
    """
    Surgically modifies the source DOCX document with proposed requirement changes
    and appends a new revision history record.
    """
    from typing import Optional
    new_rev_id = ""
    try:
        print(f"Rebuilding {doc_type} spec: {source_path} -> {output_path}")
        doc = docx.Document(source_path)
        
        # 1. Separate changes
        inserts = [c for c in changes if c["action"] == "INSERT"]
        modifies = [c for c in changes if c["action"] == "MODIFY"]
        deletes = [c for c in changes if c["action"] == "DELETE"]

        # 2. Process Deletes
        for dec in deletes:
            req_id = dec["req_id"].strip().upper()
            removed = False
            for table in doc.tables:
                for row in list(table.rows):
                    if len(row.cells) > 0 and row.cells[0].text.strip().upper() == req_id:
                        remove_row(table, row)
                        print(f"  -> Deleted row: {req_id}")
                        removed = True
                        break
                if removed:
                    break

        # 3. Process Modifies — F19: uses shared populate_row_cells helper
        for mod in modifies:
            req_id = mod["req_id"].strip().upper()
            modified = False
            for table in doc.tables:
                for row in table.rows:
                    if len(row.cells) > 0 and row.cells[0].text.strip().upper() == req_id:
                        populate_row_cells(row, doc_type, mod)
                        print(f"  -> Modified row: {req_id}")
                        modified = True
                        break
                if modified:
                    break

            if not modified:
                print(f"  -> Target row {req_id} not found for MODIFY. Falling back to INSERT.")
                inserts.append(mod)

        # 4. Process Inserts — Sort in strict numeric ID order first
        def get_numeric_sort_key(req):
            req_id = req.get("req_id", "")
            match = re.search(r'(\d+)$', req_id)
            num = int(match.group(1)) if match else 0
            return (req.get("category", ""), num, req_id)

        inserts.sort(key=get_numeric_sort_key)

        for ins in inserts:
            req_id = ins["req_id"]
            specified_table_idx = ins.get("table_index")

            best_table_idx = -1
            best_insert_after_idx = -1
            best_insert_after_number = -1

            # Check if specified table_index is valid
            if specified_table_idx is not None and isinstance(specified_table_idx, int) and 0 <= specified_table_idx < len(doc.tables):
                best_table_idx = specified_table_idx
                target_table = doc.tables[best_table_idx]
                best_insert_after_idx = len(target_table.rows) - 1

            if best_table_idx == -1:
                for table_idx, table in enumerate(doc.tables):
                    insert_after_idx = find_insert_position(table, req_id)

                    if insert_after_idx != -1:
                        candidate_number = get_requirement_row_number(table.rows[insert_after_idx])
                        if candidate_number is None:
                            candidate_number = -1

                        if candidate_number > best_insert_after_number:
                            best_insert_after_number = candidate_number
                            best_table_idx = table_idx
                            best_insert_after_idx = insert_after_idx

            if best_table_idx != -1:
                target_table = doc.tables[best_table_idx]

                # Check if target table currently contains an N/A placeholder row at row 1
                is_na_table = False
                if len(target_table.rows) > 1:
                    row1_col0 = target_table.rows[1].cells[0].text.strip().upper() if len(target_table.rows[1].cells) > 0 else ""
                    if row1_col0 in ["N/A", "NA"]:
                        is_na_table = True

                if is_na_table:
                    target_row = target_table.rows[1]
                    sync_and_clear_row_cells(target_table, target_row)
                    add_paragraph_with_run(target_row.cells[0], req_id)
                    populate_row_cells(target_row, doc_type, ins, target_table)
                    print(f"  -> Replaced N/A placeholder row in table {best_table_idx} with new requirement {req_id}")
                else:
                    new_row = insert_row_after(target_table, best_insert_after_idx)
                    sync_and_clear_row_cells(target_table, new_row)
                    print(f"  -> Inserted row after index {best_insert_after_idx} in table {best_table_idx}")
                    add_paragraph_with_run(new_row.cells[0], req_id)
                    populate_row_cells(new_row, doc_type, ins, target_table)
            else:
                # Fallback: Find target table dynamically by header content or prefix
                prefix_match = re.match(r'^([A-Z]+-[A-Z]+-)', req_id)
                if not prefix_match:
                    prefix_match = re.match(r'^([A-Z]+-)', req_id)
                prefix = prefix_match.group(1) if prefix_match else ""

                fallback_table_idx, fallback_table = find_requirement_table(doc, prefix)
                
                if fallback_table is None:
                    # Dynamically select the last non-revision requirement table in the document
                    for tidx in range(len(doc.tables) - 1, -1, -1):
                        tbl = doc.tables[tidx]
                        if len(tbl.rows) > 0:
                            hdr = " ".join(c.text.strip().lower() for c in tbl.rows[0].cells)
                            if not any(rev in hdr for rev in ['revision', 'authored on', 'reason for esign', 'reviewed on', 'approved on']):
                                fallback_table_idx = tidx
                                fallback_table = tbl
                                break
                    
                if fallback_table is not None:
                    # Check if fallback table is an N/A placeholder table
                    is_na_table = False
                    if len(fallback_table.rows) > 1:
                        row1_col0 = fallback_table.rows[1].cells[0].text.strip().upper() if len(fallback_table.rows[1].cells) > 0 else ""
                        if row1_col0 in ["N/A", "NA"]:
                            is_na_table = True

                    if is_na_table:
                        target_row = fallback_table.rows[1]
                        sync_and_clear_row_cells(fallback_table, target_row)
                        add_paragraph_with_run(target_row.cells[0], req_id)
                        populate_row_cells(target_row, doc_type, ins, fallback_table)
                        print(f"  -> Replaced N/A placeholder row in table {fallback_table_idx} with new requirement {req_id}")
                    else:
                        last_row_idx = len(fallback_table.rows) - 1
                        new_row = insert_row_after(fallback_table, last_row_idx)
                        sync_and_clear_row_cells(fallback_table, new_row)
                        add_paragraph_with_run(new_row.cells[0], req_id)
                        populate_row_cells(new_row, doc_type, ins, fallback_table)
                        print(f"  -> Appended row to table {fallback_table_idx}")

        # 5. Insert Latest Revision History Record
        rev_table = find_revision_table(doc)
        if rev_table is not None:
            versions = []
            version_pattern = r'(\d+)$'
            
            # Read version numbers across all rows in the revision table
            for row in rev_table.rows[1:]:
                if len(row.cells) > 0:
                    text = row.cells[0].text.strip().replace('\n', '').replace('\r', '').replace(' ', '')
                    match = re.search(version_pattern, text)
                    if match:
                        versions.append((int(match.group(1)), text))
            
            # Find the highest version to generate the next version ID
            if versions:
                max_val, max_rev_id = max(versions, key=lambda x: x[0])
                digit_match = re.search(version_pattern, max_rev_id)
                digits = digit_match.group(1)
                inc_digits = f"{max_val + 1:0{len(digits)}d}"
                new_rev_id = max_rev_id[:-len(digits)] + inc_digits
            else:
                # Fallback if no version ID is found in the table
                new_rev_id = f"{doc_type}-G4.NPSS.CU.0069-0001.0002"
            
            # Determine ordering: top-down (chronological) vs bottom-up (reverse chronological)
            is_top_down = False
            if len(versions) >= 2:
                first_val = versions[0][0]
                last_val = versions[-1][0]
                if first_val < last_val:
                    is_top_down = True
            
            # Insert row in correct location
            if is_top_down:
                # Top-down: append at the very bottom
                template_row_idx = len(rev_table.rows) - 1
                new_rev_row = insert_row_after(rev_table, template_row_idx, template_row_idx=template_row_idx)
                print(f"  -> Appended latest revision at the bottom of FIRST revision table (Row {len(rev_table.rows)-1}): {new_rev_id}")
            else:
                # Bottom-up: insert at top (index 0, immediately after header row)
                template_row_idx = 1 if len(rev_table.rows) > 1 else 0
                new_rev_row = insert_row_after(rev_table, 0, template_row_idx=template_row_idx)
                print(f"  -> Inserted latest revision at top of FIRST revision table (Row 1): {new_rev_id}")
            
            # Format revision bullet-tabs for FRS specifically
            # Call Claude for a professional compliance-standard FDA/GAMP reason for change if it is generic/default
            if not revision_reason or revision_reason == "Updated as per user story modifications.":
                try:
                    from services.claude import ClaudeClient
                    claude = ClaudeClient()
                    
                    new_reqs = [c for c in changes if c.get("action") == "INSERT"]
                    mod_reqs = [c for c in changes if c.get("action") == "MODIFY"]
                    del_reqs = [c for c in changes if c.get("action") == "DELETE"]
                    
                    if new_reqs or mod_reqs or del_reqs:
                        system_prompt = """<role>
You are an expert systems compliance editor writing a very concise "Reason for Change" for GAMP 5 / FDA regulated specification documents.
</role>

<rules>
<rule name="format">Provide an extremely concise summary of the business objective or release context. You MUST use exactly 6 to 7 words.</rule>
<rule name="tone">Use compliance-professional tone (e.g. "Updated as per release 8 requirements" or "Added QMS training entity report validation").</rule>
<rule name="no_extras">Do not include headers, intros, quotes, markdown, or commentary. Respond ONLY with the 6-7 words text.</rule>
</rules>"""

                        user_prompt = f"""<change_data>
Document Type: {doc_type}
New Requirements: {[n.get('req_id') or 'NEW_ID' for n in new_reqs]}
Modified Requirements: {[m.get('req_id') or 'MOD_ID' for m in mod_reqs]}
Deleted Requirements: {[d.get('req_id') or 'DEL_ID' for d in del_reqs]}
</change_data>

Generate the compliance-standard Reason for Change (exactly 6-7 words).
"""
                        revision_reason = claude.query(system_prompt, user_prompt).strip().strip('"').strip("'")
                    else:
                        revision_reason = "No changes."
                except Exception as e:
                    print(f"Error generating reason with Claude in writer: {e}")
                    revision_reason = f"Updated as per release 8 requirements."

            cleaned_log = revision_log
            if doc_type == 'FRS':
                if not cleaned_log.startswith('•\t') and not cleaned_log.startswith('-'):
                    cleaned_log = '•\t' + cleaned_log.replace('\n', '\n•\t')

            add_paragraph_with_run(new_rev_row.cells[0], new_rev_id)
            add_paragraph_with_run(new_rev_row.cells[1], "NEW") # Kept as NEW per user request
            add_paragraph_with_run(new_rev_row.cells[2], cleaned_log)
            add_paragraph_with_run(new_rev_row.cells[3], revision_reason)
            
        if metadata_out is not None:
            metadata_out["new_rev_id"] = new_rev_id
            print(f"  -> metadata_out[new_rev_id] = '{new_rev_id}'")

        # Save document
        doc.save(output_path)
        print(f"Rebuild completed successfully for {doc_type}.")
        return True

    except Exception as e:
        print(f"Error rebuilding document: {e}")
        import traceback
        traceback.print_exc()
        return False